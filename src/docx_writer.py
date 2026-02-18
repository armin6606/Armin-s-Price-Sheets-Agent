"""DOCX writer - updates Word templates with price sheet data.

Formatting rules applied to every data row written:
  - All cells: vertically centered + horizontally centered
  - Notes contains "sold"              → entire row font = RED
  - Notes contains "upgraded flooring" → NOTES cell only: font = WHITE, cell fill = PURPLE
                                         (rest of row stays normal)
"""

import io
import logging
from typing import Optional

from docx import Document
from docx.table import Table
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

from .utils import (
    build_header_map,
    validate_headers,
    format_price,
    parse_ready_by,
    normalize_for_compare,
)
from .control_parser import ControlRow
from .locator import TableMatch

logger = logging.getLogger("price_sheet_bot.docx_writer")

# ── Formatting constants ──
COLOR_RED    = RGBColor(0xFF, 0x00, 0x00)
COLOR_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_BLACK  = RGBColor(0x00, 0x00, 0x00)

# Purple: a rich purple (#7030A0 is Word's standard purple)
HEX_PURPLE   = "7030A0"
HEX_RED_FILL = "FF0000"   # used if you ever want red fill (not needed currently)


class DocxWriteResult:
    """Result of a DOCX write operation."""
    def __init__(self):
        self.action: str = ""  # "appended", "updated_blanks", "overwritten", "skipped"
        self.row_index: int = -1
        self.details: str = ""
        self.error: str = ""


# ── Low-level cell helpers ──

def _get_cell_texts(table: Table, row_index: int) -> list:
    """Get text of all cells in a table row."""
    if row_index >= len(table.rows):
        return []
    return [cell.text.strip() for cell in table.rows[row_index].cells]


def _set_cell_text(table: Table, row_index: int, col_index: int, value: str):
    """Set the text of a cell, preserving first run formatting if possible."""
    if row_index >= len(table.rows):
        return
    cell = table.rows[row_index].cells[col_index]
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = value
        for run in cell.paragraphs[0].runs[1:]:
            run.text = ""
        for para in cell.paragraphs[1:]:
            for run in para.runs:
                run.text = ""
    else:
        cell.text = value


def _is_row_blank(table: Table, row_index: int, header_map: dict) -> bool:
    """Check if a row is blank (all mapped columns are empty/whitespace)."""
    if row_index >= len(table.rows):
        return True
    cells = _get_cell_texts(table, row_index)
    for col_idx in header_map.values():
        if col_idx < len(cells) and cells[col_idx].strip():
            return False
    return True


def _find_existing_site_rows(table: Table, header_map: dict, homesite: str,
                              data_start: int) -> list:
    """Find row indices in the table where SITE matches the homesite."""
    site_col = header_map.get("SITE")
    if site_col is None:
        return []
    h_norm = normalize_for_compare(homesite)
    matches = []
    for r_idx in range(data_start, len(table.rows)):
        cells = _get_cell_texts(table, r_idx)
        if site_col < len(cells):
            if normalize_for_compare(cells[site_col]) == h_norm:
                matches.append(r_idx)
    return matches


def _find_next_blank_row(table: Table, header_map: dict, data_start: int) -> int:
    """Find the next blank row starting from data_start. Returns -1 if none."""
    for r_idx in range(data_start, len(table.rows)):
        if _is_row_blank(table, r_idx, header_map):
            return r_idx
    return -1


def _add_row_to_table(table: Table) -> int:
    """Add a new blank row at the end of a table, cloning the structure of the last row.

    Copies the last row's XML structure (cell count, widths, borders) but
    clears all text content.  Returns the index of the newly added row.
    """
    import copy
    from lxml import etree

    last_row = table.rows[-1]
    # Deep copy the last row's XML element
    new_tr = copy.deepcopy(last_row._tr)

    # Clear all text in every cell of the new row
    for tc in new_tr.findall(qn("w:tc")):
        for p in tc.findall(qn("w:p")):
            # Keep paragraph properties (alignment, etc.) but clear runs
            for r in p.findall(qn("w:r")):
                p.remove(r)
            # Also clear any bare text nodes
            for t in p.findall(qn("w:t")):
                p.remove(t)

    # Append the new row to the table's XML
    table._tbl.append(new_tr)
    new_index = len(table.rows) - 1
    logger.info("Added new row to table (now %d rows), new row index = %d",
                len(table.rows), new_index)
    return new_index


# ── Formatting helpers ──

def _set_cell_shading(cell, hex_color: str):
    """Set the background fill colour of a table cell using OOXML shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing shading element
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper())
    tcPr.append(shd)


def _clear_cell_shading(cell):
    """Remove all background fill/shading from a table cell (reset to no fill)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)


def _set_cell_font_color(cell, rgb: RGBColor):
    """Set font colour for every run in every paragraph of a cell."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = rgb
        # If paragraph has no runs but has text, we need to handle it via XML
        if not para.runs and para.text.strip():
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            color_el = OxmlElement("w:color")
            color_el.set(qn("w:val"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
            rPr.append(color_el)
            r.append(rPr)
            t = OxmlElement("w:t")
            t.text = para.text
            r.append(t)
            para._p.clear()
            para._p.append(rPr)
            para._p.append(r)


def _set_row_alignment(table: Table, row_index: int):
    """Centre every cell in a row: vertically (middle) + horizontally (center)."""
    if row_index >= len(table.rows):
        return
    row = table.rows[row_index]
    for cell in row.cells:
        # Vertical alignment - middle
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Horizontal alignment - center for every paragraph
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _determine_row_style(notes: str) -> str:
    """Return the formatting style to apply based on the notes field.

    Returns: 'sold', 'upgraded_flooring', or 'normal'
    """
    notes_lower = notes.strip().lower()
    if "sold" in notes_lower:
        return "sold"
    if "upgraded flooring" in notes_lower:
        return "upgraded_flooring"
    return "normal"


def _apply_row_formatting(table: Table, row_index: int, notes: str,
                          header_map: dict = None):
    """Apply colour formatting + centre alignment to a data row.

    Rules:
      - sold              → red font on ALL cells in the row
      - upgraded flooring → ONLY the NOTES cell: white font + purple background
                            rest of row: normal (black font, no background)
      - normal            → black font + no cell background (reset to default)
      - always            → vertically and horizontally centred
    """
    if row_index >= len(table.rows):
        return

    style = _determine_row_style(notes)
    row = table.rows[row_index]

    # Identify the NOTES column index so we can treat it specially
    notes_col = header_map.get("NOTES", -1) if header_map else -1

    for col_idx, cell in enumerate(row.cells):
        # ── Text colour + shading ──
        if style == "sold":
            # Sold: entire row gets red font, no background
            _set_cell_font_color(cell, COLOR_RED)
            _clear_cell_shading(cell)
        elif style == "upgraded_flooring":
            if col_idx == notes_col:
                # Only the NOTES cell gets white font + purple background
                _set_cell_font_color(cell, COLOR_WHITE)
                _set_cell_shading(cell, HEX_PURPLE)
            else:
                # All other cells: normal formatting
                _set_cell_font_color(cell, COLOR_BLACK)
                _clear_cell_shading(cell)
        else:
            # Normal: reset font to black and remove any background shading
            _set_cell_font_color(cell, COLOR_BLACK)
            _clear_cell_shading(cell)

        # ── Alignment ──
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── Main write helpers ──

def _write_row(table: Table, row_index: int, header_map: dict, control: ControlRow):
    """Write all fields to a row, then apply formatting."""
    if "SITE" in header_map:
        _set_cell_text(table, row_index, header_map["SITE"], str(control.homesite))
    if "PRICE" in header_map:
        _set_cell_text(table, row_index, header_map["PRICE"], format_price(control.price))
    if "ADDRESS" in header_map:
        _set_cell_text(table, row_index, header_map["ADDRESS"], control.address)
    if "READY BY" in header_map:
        _set_cell_text(table, row_index, header_map["READY BY"], parse_ready_by(control.ready_by))
    if "NOTES" in header_map:
        _set_cell_text(table, row_index, header_map["NOTES"], control.notes)

    # Apply formatting after text is set
    _apply_row_formatting(table, row_index, control.notes, header_map=header_map)


def _fill_blank_cells(table: Table, row_index: int, header_map: dict,
                       control: ControlRow, allow_price_update: bool):
    """Fill only blank cells in an existing row, then refresh formatting."""
    cells = _get_cell_texts(table, row_index)

    def _is_blank(col_name):
        col_idx = header_map.get(col_name)
        if col_idx is None:
            return False
        return col_idx >= len(cells) or not cells[col_idx].strip()

    if "ADDRESS" in header_map and _is_blank("ADDRESS"):
        _set_cell_text(table, row_index, header_map["ADDRESS"], control.address)
    if "READY BY" in header_map and _is_blank("READY BY"):
        _set_cell_text(table, row_index, header_map["READY BY"], parse_ready_by(control.ready_by))
    if "NOTES" in header_map and _is_blank("NOTES"):
        _set_cell_text(table, row_index, header_map["NOTES"], control.notes)
    if allow_price_update and "PRICE" in header_map and _is_blank("PRICE"):
        _set_cell_text(table, row_index, header_map["PRICE"], format_price(control.price))

    # Always re-apply formatting so new notes value is reflected
    _apply_row_formatting(table, row_index, control.notes, header_map=header_map)


# ── Public API ──

def write_to_template(
    doc_bytes: bytes,
    table_match_info: dict,
    control_row: ControlRow,
    overwrite_existing: bool = False,
    update_only_blank_cells: bool = True,
    allow_price_update: bool = False,
    strict_mode: bool = True,
    remove_invisible_code: bool = True,
    header_row_1based: int = 2,
) -> tuple:
    """Write control row data into the template DOCX.

    Args:
        doc_bytes: Raw DOCX bytes
        table_match_info: Dict with invisible_code (and optionally table_index)
        control_row: The CONTROL row data to write
        overwrite_existing: If True, overwrite existing site rows
        update_only_blank_cells: If True and site exists, only fill blank cells
        allow_price_update: If True, allow price update when filling blanks
        strict_mode: If True, fail on duplicate sites in same table
        remove_invisible_code: If True, remove invisible code after finding table
        header_row_1based: Header row number (1-based), default 2

    Returns:
        (modified_doc_bytes, DocxWriteResult)
    """
    from .locator import find_table_by_invisible_code, remove_invisible_code as do_remove

    result = DocxWriteResult()
    doc = Document(io.BytesIO(doc_bytes))

    # Find target table
    invisible_code = table_match_info["invisible_code"]
    match = find_table_by_invisible_code(doc, invisible_code)
    table = match.table

    # Header row (0-based index)
    header_idx = header_row_1based - 1
    if header_idx >= len(table.rows):
        result.error = (
            f"Header row {header_row_1based} is beyond table size ({len(table.rows)} rows)."
        )
        return None, result

    # Build header map
    header_cells = _get_cell_texts(table, header_idx)
    header_map = build_header_map(header_cells)

    missing = validate_headers(header_map, strict=strict_mode)
    if missing:
        result.error = f"Missing required headers: {missing}. Found: {list(header_map.keys())}"
        return None, result

    data_start = header_idx + 1

    # Check for existing site rows within THIS table only
    existing_rows = _find_existing_site_rows(
        table, header_map, control_row.homesite, data_start
    )

    if len(existing_rows) > 1 and strict_mode:
        result.error = (
            f"Duplicate site '{control_row.homesite}' found at rows "
            f"{existing_rows} within the same target table. "
            f"strict_mode requires unique sites per table."
        )
        result.action = "duplicate_site_rows_in_table"
        return None, result

    if existing_rows:
        target_row = existing_rows[0]
        if overwrite_existing:
            _write_row(table, target_row, header_map, control_row)
            result.action = "overwritten"
            result.row_index = target_row
            result.details = f"Overwrote existing site at row {target_row + 1}"
        elif update_only_blank_cells:
            _fill_blank_cells(
                table, target_row, header_map, control_row, allow_price_update
            )
            result.action = "updated_blanks"
            result.row_index = target_row
            result.details = f"Updated blank cells at row {target_row + 1}"
        else:
            result.action = "duplicate_site_same_table"
            result.error = (
                f"Site '{control_row.homesite}' already exists at row "
                f"{target_row + 1}. Skipping."
            )
            return None, result
    else:
        blank_row = _find_next_blank_row(table, header_map, data_start)
        if blank_row < 0:
            # No blank rows available — add a new row to the table
            logger.info(
                "No blank row in table[%d] (%d rows). Adding a new row for site=%s.",
                match.table_index, len(table.rows), control_row.homesite,
            )
            blank_row = _add_row_to_table(table)

        _write_row(table, blank_row, header_map, control_row)
        result.action = "appended"
        result.row_index = blank_row
        result.details = f"Appended to row {blank_row + 1}"

    # Remove invisible code if configured
    if remove_invisible_code:
        do_remove(match, invisible_code)

    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    modified_bytes = output.getvalue()

    logger.info(
        "DOCX write: %s for site=%s in table[%d] at row %d",
        result.action, control_row.homesite, match.table_index, result.row_index + 1,
    )
    return modified_bytes, result
