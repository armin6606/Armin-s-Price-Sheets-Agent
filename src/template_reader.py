"""Template reader - extracts existing homesite data from Word templates.

Opens each template DOCX from Drive, finds the data table (via invisible code),
reads all non-blank data rows, and returns them as a list of dicts that can be
synced back to the CONTROL tab.
"""

import io
import logging
from typing import List, Optional

from docx import Document

from .locator import find_table_by_invisible_code
from .utils import build_header_map, normalize_for_compare

logger = logging.getLogger("price_sheet_bot.template_reader")


def read_template_data(
    doc_bytes: bytes,
    invisible_code: str,
    community: str,
    floorplan: str,
    header_row_1based: int = 2,
) -> List[dict]:
    """Read all existing data rows from a template's target table.

    Args:
        doc_bytes: Raw DOCX bytes.
        invisible_code: The invisible code string to locate the table.
        community: Community name (from MAPPING).
        floorplan: Floorplan name (from MAPPING).
        header_row_1based: Header row (1-based), default 2.

    Returns:
        List of dicts, each with keys: community, homesite, floorplan, price,
        address, ready_by, notes.  Only returns rows where at least SITE is
        non-blank.
    """
    doc = Document(io.BytesIO(doc_bytes))

    try:
        match = find_table_by_invisible_code(doc, invisible_code)
    except ValueError as e:
        logger.warning("Cannot find table in template: %s", e)
        return []

    table = match.table
    header_idx = header_row_1based - 1

    if header_idx >= len(table.rows):
        logger.warning("Header row %d beyond table size (%d rows).", header_row_1based, len(table.rows))
        return []

    # Build header map
    header_cells = [cell.text.strip() for cell in table.rows[header_idx].cells]
    hmap = build_header_map(header_cells)

    # Need at least SITE column
    if "SITE" not in hmap:
        logger.warning("No SITE column found in template headers: %s", header_cells)
        return []

    site_col = hmap["SITE"]
    price_col = hmap.get("PRICE")
    addr_col = hmap.get("ADDRESS")
    rb_col = hmap.get("READY BY")
    notes_col = hmap.get("NOTES")

    # Read data rows (start after header)
    data_start = header_idx + 1
    rows_found = []

    for r_idx in range(data_start, len(table.rows)):
        cells = [cell.text.strip() for cell in table.rows[r_idx].cells]

        # Get the SITE value
        site_val = cells[site_col] if site_col < len(cells) else ""
        if not site_val.strip():
            continue  # Skip blank rows

        row_data = {
            "community": community,
            "homesite": site_val.strip(),
            "floorplan": floorplan,
            "price": "",
            "address": "",
            "ready_by": "",
            "notes": "",
        }

        if price_col is not None and price_col < len(cells):
            row_data["price"] = cells[price_col].strip()
        if addr_col is not None and addr_col < len(cells):
            row_data["address"] = cells[addr_col].strip()
        if rb_col is not None and rb_col < len(cells):
            row_data["ready_by"] = cells[rb_col].strip()
        if notes_col is not None and notes_col < len(cells):
            row_data["notes"] = cells[notes_col].strip()

        rows_found.append(row_data)

    logger.info(
        "Read %d existing rows from template (community=%s, floorplan=%s)",
        len(rows_found), community, floorplan,
    )
    return rows_found


def read_final_docx_data(
    doc_bytes: bytes,
    community: str,
    floorplan: str,
    header_row_1based: int = 2,
) -> List[dict]:
    """Read data from a FINAL price sheet DOCX that has NO invisible code.

    Scans ALL tables for those whose header row matches the expected pattern
    (SITE, PRICE, etc.) and reads data from ALL of them.  Since invisible
    codes are removed from final files, we cannot distinguish which table
    belongs to which floorplan — so we return ALL data rows from ALL
    matching tables.  The caller can match by homesite.

    Args:
        doc_bytes: Raw DOCX bytes of the final price sheet.
        community: Community name (from MAPPING).
        floorplan: Floorplan name (from MAPPING).
        header_row_1based: Header row (1-based), default 2.

    Returns:
        List of dicts with keys: community, homesite, floorplan, price,
        address, ready_by, notes.
    """
    doc = Document(io.BytesIO(doc_bytes))
    header_idx = header_row_1based - 1
    rows_found = []

    for table in doc.tables:
        if header_idx >= len(table.rows):
            continue

        header_cells = [cell.text.strip() for cell in table.rows[header_idx].cells]
        hmap = build_header_map(header_cells)

        # Must have at least SITE column to be a valid data table
        if "SITE" not in hmap:
            continue

        site_col = hmap["SITE"]
        price_col = hmap.get("PRICE")
        addr_col = hmap.get("ADDRESS")
        rb_col = hmap.get("READY BY")
        notes_col = hmap.get("NOTES")

        data_start = header_idx + 1
        for r_idx in range(data_start, len(table.rows)):
            cells = [cell.text.strip() for cell in table.rows[r_idx].cells]

            site_val = cells[site_col] if site_col < len(cells) else ""
            if not site_val.strip():
                continue

            row_data = {
                "community": community,
                "homesite": site_val.strip(),
                "floorplan": floorplan,
                "price": "",
                "address": "",
                "ready_by": "",
                "notes": "",
            }

            if price_col is not None and price_col < len(cells):
                row_data["price"] = cells[price_col].strip()
            if addr_col is not None and addr_col < len(cells):
                row_data["address"] = cells[addr_col].strip()
            if rb_col is not None and rb_col < len(cells):
                row_data["ready_by"] = cells[rb_col].strip()
            if notes_col is not None and notes_col < len(cells):
                row_data["notes"] = cells[notes_col].strip()

            rows_found.append(row_data)

    logger.info(
        "Read %d rows from final DOCX (community=%s, floorplan=%s) via header scan",
        len(rows_found), community, floorplan,
    )
    return rows_found


def read_all_templates(
    drive_client,
    templates_folder_id: str,
    mapping_rows: list,
    download_cache_dir: str,
) -> List[dict]:
    """Read data from ALL templates listed in MAPPING.

    Downloads each template, reads its data rows, returns a combined list
    of all homesite dicts across all templates.
    """
    import os
    all_rows = []
    seen_templates = set()

    for mrow in mapping_rows:
        # Avoid re-downloading the same template file multiple times
        # (some templates serve multiple floorplans)
        template_key = mrow.file_name
        if template_key in seen_templates:
            # Still need to read for this floorplan, but use cached bytes
            cache_path = os.path.join(download_cache_dir, mrow.file_name)
            if os.path.exists(cache_path):
                with open(cache_path, "rb") as f:
                    doc_bytes = f.read()
                rows = read_template_data(
                    doc_bytes, mrow.invisible_code,
                    mrow.community, mrow.floorplan,
                    header_row_1based=mrow.header_row,
                )
                all_rows.extend(rows)
            continue

        seen_templates.add(template_key)

        # Find and download template
        tf = drive_client.find_file_by_name(templates_folder_id, mrow.file_name)
        if not tf:
            logger.warning("Template '%s' not found in Drive, skipping.", mrow.file_name)
            continue

        cache_path = os.path.join(download_cache_dir, mrow.file_name)
        try:
            drive_client.download_file(tf["id"], cache_path)
        except Exception as e:
            logger.warning("Failed to download template '%s': %s", mrow.file_name, e)
            continue

        with open(cache_path, "rb") as f:
            doc_bytes = f.read()

        rows = read_template_data(
            doc_bytes, mrow.invisible_code,
            mrow.community, mrow.floorplan,
            header_row_1based=mrow.header_row,
        )
        all_rows.extend(rows)

    logger.info("Total rows read from all templates: %d", len(all_rows))
    return all_rows
