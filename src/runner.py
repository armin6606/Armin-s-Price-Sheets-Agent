"""Runner - main processing engine for Price Sheet Bot.

New-release PDFs contain MULTIPLE homesites per file.  The runner:
  1. Downloads each PDF from the New Releases Drive folder.
  2. Extracts homesite rows (HS #, Plan, prices, COE, etc.) via pdf_parser.
  3. Groups homesites by (community, plan) -> MAPPING row -> template.
  4. For each template, writes ALL matching homesite rows into the DOCX.
  5. Converts to PDF and uploads both DOCX + PDF to Final Price Sheets.

The CONTROL tab is used as an *optional* supplement: if a CONTROL row
matches (community, homesite, floorplan), its address / notes override
blanks.  Prices always come from the release PDF.
"""

import json
import logging
import os
import platform
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional, List

from .config import Config
from .sheets import SheetsClient
from .drive_client import DriveClient
from .control_parser import parse_control_tab, find_control_row, ControlRow
from .mapping_parser import parse_mapping_tab, find_mapping_row
from .docx_writer import write_to_template
from .pdf_export import export_to_pdf
from .sop_resolver import resolve_address
from .locator import find_table_by_invisible_code, scan_template_for_markers
from .pdf_parser import parse_release_pdf, parse_release_filename, ReleaseHomesite
from .utils import (
    parse_pdf_filename, compute_hash, normalize_for_compare,
    format_price, parse_ready_by,
)
from .logging_setup import log_event

logger = logging.getLogger("price_sheet_bot.runner")


# ── Write Lock ──

LOCK_FILE = "./cache/.process_lock"
LOCK_TIMEOUT_SECONDS = 600  # 10 minutes


def acquire_lock(force: bool = False) -> bool:
    """Acquire the process lock. Returns True if acquired."""
    os.makedirs(os.path.dirname(LOCK_FILE), exist_ok=True)

    if os.path.exists(LOCK_FILE):
        try:
            with open(LOCK_FILE, "r") as f:
                lock_data = json.load(f)
            lock_time = datetime.fromisoformat(lock_data["timestamp"])
            age = (datetime.now(timezone.utc) - lock_time).total_seconds()

            if age < LOCK_TIMEOUT_SECONDS and not force:
                logger.error(
                    "Process lock held by PID %s since %s (%.0fs ago). "
                    "Use --force-lock-reset to override.",
                    lock_data.get("pid"), lock_data.get("timestamp"), age,
                )
                return False
            else:
                if age >= LOCK_TIMEOUT_SECONDS:
                    logger.warning("Stale lock detected (%.0fs old). Overriding.", age)
        except (json.JSONDecodeError, KeyError, ValueError):
            logger.warning("Corrupt lock file. Overriding.")

    lock_data = {
        "pid": os.getpid(),
        "hostname": platform.node(),
        "timestamp": datetime.now(timezone.utc).isoformat(),
    }
    with open(LOCK_FILE, "w") as f:
        json.dump(lock_data, f)
    return True


def release_lock():
    """Release the process lock."""
    if os.path.exists(LOCK_FILE):
        os.remove(LOCK_FILE)


# ── Processed Manifest ──

def load_manifest(path: str) -> dict:
    """Load the processed manifest from disk."""
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def save_manifest(path: str, manifest: dict):
    """Save the processed manifest to disk."""
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(manifest, f, indent=2, default=str)


def is_already_processed(pdf_file: dict, manifest: dict) -> bool:
    """Check if a PDF has already been processed."""
    props = pdf_file.get("appProperties", {})
    if props.get("processed") == "true":
        return True
    if pdf_file["id"] in manifest:
        return True
    return False


# ── Template Certification ──

CERT_FILE = "./cache/certified_templates.json"


def load_certifications() -> dict:
    if os.path.exists(CERT_FILE):
        with open(CERT_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def save_certifications(certs: dict):
    os.makedirs(os.path.dirname(CERT_FILE), exist_ok=True)
    with open(CERT_FILE, "w", encoding="utf-8") as f:
        json.dump(certs, f, indent=2, default=str)


def is_template_certified(file_id: str, modified_time: str, certs: dict) -> bool:
    if file_id not in certs:
        return False
    return certs[file_id].get("modifiedTime") == modified_time


# ── Quarantine ──

def quarantine_pdf(drive_client: DriveClient, pdf_file: dict,
                    final_folder_id: str, reason: str,
                    quarantine_folder_name: str = "Quarantine"):
    """Move a PDF to the quarantine subfolder with a reason tag."""
    q_folder_id = drive_client.ensure_subfolder(final_folder_id, quarantine_folder_name)
    try:
        drive_client.move_file(pdf_file["id"], q_folder_id)
        drive_client.set_app_properties(pdf_file["id"], {
            "quarantined": "true",
            "quarantine_reason": reason,
            "quarantined_at": datetime.now(timezone.utc).isoformat(),
        })
        logger.warning("Quarantined PDF '%s' (id=%s): %s", pdf_file["name"], pdf_file["id"], reason)
    except Exception as e:
        logger.error("Failed to quarantine PDF '%s': %s", pdf_file["name"], e)


# ── Helper: build ControlRow from PDF homesite data ──

def _build_control_row_from_pdf(
    hs: ReleaseHomesite,
    control_rows: list,
    drive_client: DriveClient = None,
    sop_folder_id: str = None,
) -> ControlRow:
    """Build a ControlRow from PDF-extracted homesite data.

    Uses the net_price from the PDF as the price.
    Uses the COE date from the PDF as ready_by.
    If a matching CONTROL row exists, pulls address/notes from it.
    """
    community = hs.community
    homesite = hs.homesite
    floorplan = hs.plan

    # Try to find a matching CONTROL row for supplemental data
    ctrl = find_control_row(control_rows, community, homesite, floorplan)

    # Address: prefer CONTROL, then SOP, then blank
    address = ""
    if ctrl and ctrl.address.strip():
        address = ctrl.address
    elif drive_client and sop_folder_id:
        try:
            addr = resolve_address(drive_client, sop_folder_id, community, homesite, floorplan)
            if addr:
                address = addr
        except Exception:
            pass

    # Notes: prefer CONTROL, then blank
    notes = ""
    if ctrl and ctrl.notes.strip():
        notes = ctrl.notes

    # Price: ALWAYS from PDF (total_released_price)
    price = hs.total_released_price if hs.total_released_price else hs.base_price

    # Ready-by / COE: from PDF homesite row COE date, fall back to header-level COE
    # move_in stores the raw date string from the PDF (e.g. "April, 2026")
    # ready_by stores the same value — parse_ready_by() in docx_writer formats it
    raw_coe = hs.coe_date if hs.coe_date else hs.default_coe
    ready_by = raw_coe

    # If a CONTROL row has a ready_by/move_in date already set, prefer it
    # (allows manual overrides in the sheet to survive re-processing)
    if ctrl and ctrl.ready_by.strip():
        ready_by = ctrl.ready_by
        raw_coe = ctrl.move_in or ctrl.ready_by

    return ControlRow(
        enabled=True,
        community=community,
        homesite=homesite,
        floorplan=floorplan,
        price=price,
        address=address,
        ready_by=ready_by,
        move_in=raw_coe,
        notes=notes,
        row_index=-1,  # Not from a sheet row
    )


# ── Main Processing (New PDF-driven flow) ──

def process_release_pdf(
    cfg: Config,
    drive_client: DriveClient,
    sheets_client: SheetsClient,
    control_rows: list,
    mapping_rows: list,
    pdf_file: dict,
    manifest: dict,
    certs: dict,
    overwrite_existing_override: bool = False,
    homesite_filter: str = None,
    floorplan_filter: str = None,
) -> dict:
    """Process a single release PDF containing multiple homesites.

    1. Download the PDF from Drive.
    2. Parse it to extract metadata + homesite rows.
    3. For each homesite, find the MAPPING row -> template.
    4. Group homesites by template (one template may get multiple rows).
    5. For each template: download, write all rows, upload DOCX + PDF.

    Returns result dict with status, details, per-homesite results, etc.
    """
    pdf_name = pdf_file["name"]
    pdf_id = pdf_file["id"]
    result = {
        "pdf_name": pdf_name, "pdf_id": pdf_id,
        "status": "unknown", "details": "",
        "homesite_results": [], "output_ids": {},
    }

    log_event(logger, "INFO", "Processing release PDF", pdf_name=pdf_name, pdf_id=pdf_id)

    # Step 1: Download the PDF to local cache
    cfg.ensure_cache_dirs()
    local_pdf_path = os.path.join(cfg.drive.download_cache_dir, pdf_name)
    try:
        drive_client.download_file(pdf_id, local_pdf_path)
    except Exception as e:
        result["status"] = "error"
        result["details"] = f"Failed to download PDF: {e}"
        quarantine_pdf(drive_client, pdf_file, cfg.drive.final_price_sheets_folder_id,
                       "download_failed", cfg.pdf.quarantine_folder_name)
        return result

    # Step 2: Parse the PDF content
    parsed_pdf = parse_release_pdf(local_pdf_path)

    if parsed_pdf.errors:
        for err in parsed_pdf.errors:
            logger.warning("PDF parse warning for '%s': %s", pdf_name, err)

    if not parsed_pdf.homesites:
        result["status"] = "error"
        result["details"] = f"No homesite rows found in PDF. Errors: {parsed_pdf.errors}"
        quarantine_pdf(drive_client, pdf_file, cfg.drive.final_price_sheets_folder_id,
                       "no_homesites_in_pdf", cfg.pdf.quarantine_folder_name)
        return result

    community = parsed_pdf.meta.community
    phase = parsed_pdf.meta.phase
    logger.info(
        "PDF '%s': community=%s phase=%s, %d homesites found.",
        pdf_name, community, phase, len(parsed_pdf.homesites),
    )

    # Step 3: Filter homesites if CLI filters provided
    homesites = parsed_pdf.homesites
    if homesite_filter:
        h_norm = normalize_for_compare(homesite_filter)
        homesites = [hs for hs in homesites if normalize_for_compare(hs.homesite) == h_norm]
    if floorplan_filter:
        f_norm = normalize_for_compare(floorplan_filter)
        homesites = [hs for hs in homesites if normalize_for_compare(hs.plan) == f_norm]

    if not homesites:
        result["status"] = "skipped"
        result["details"] = "All homesites filtered out by CLI filters."
        return result

    # Step 4: Group homesites by (community, plan) -> MAPPING row -> template
    #   Key: (template_file_name) -> list of (homesite, mapping_row)
    template_groups = {}
    errors_per_hs = []

    for hs in homesites:
        mrow = find_mapping_row(mapping_rows, community, hs.plan)
        if not mrow:
            msg = f"No MAPPING row for ({community}, {hs.plan}) - HS #{hs.homesite} skipped"
            logger.warning(msg)
            errors_per_hs.append({"homesite": hs.homesite, "plan": hs.plan, "error": msg})
            continue

        key = mrow.file_name
        if key not in template_groups:
            template_groups[key] = {"mapping_row": mrow, "homesites": []}
        template_groups[key]["homesites"].append(hs)

    if not template_groups:
        result["status"] = "error"
        result["details"] = "No MAPPING matches for any homesite in this PDF."
        result["homesite_results"] = errors_per_hs
        quarantine_pdf(drive_client, pdf_file, cfg.drive.final_price_sheets_folder_id,
                       "no_mapping_matches", cfg.pdf.quarantine_folder_name)
        return result

    # Step 5: For each template, download -> write all rows -> upload
    all_output_ids = {}
    success_count = 0
    error_count = 0

    for template_name, group in template_groups.items():
        mrow = group["mapping_row"]
        hs_list = group["homesites"]

        logger.info(
            "Template '%s': writing %d homesite(s) for community=%s",
            template_name, len(hs_list), community,
        )

        # Find template in Drive
        template_file = drive_client.find_file_by_name(cfg.drive.templates_folder_id, template_name)
        if not template_file:
            msg = f"Template '{template_name}' not found in Drive."
            logger.error(msg)
            for hs in hs_list:
                errors_per_hs.append({"homesite": hs.homesite, "plan": hs.plan, "error": msg})
                error_count += 1
            continue

        # Check certification
        if not cfg.app.allow_uncertified_templates:
            if not is_template_certified(template_file["id"], template_file.get("modifiedTime", ""), certs):
                msg = (
                    f"Template '{template_name}' is not certified. "
                    f"Run --certify-template --community {community} --floorplan {hs_list[0].plan}"
                )
                logger.error(msg)
                for hs in hs_list:
                    errors_per_hs.append({"homesite": hs.homesite, "plan": hs.plan, "error": msg})
                    error_count += 1
                continue

        # Download template
        cache_path = os.path.join(cfg.drive.download_cache_dir, template_name)
        drive_client.download_file(template_file["id"], cache_path)
        with open(cache_path, "rb") as f:
            template_bytes = f.read()

        # Write each homesite row into the template sequentially
        current_bytes = template_bytes
        write_actions = []

        for hs in hs_list:
            # Build a ControlRow from PDF data (+ optional CONTROL supplement)
            control_row = _build_control_row_from_pdf(
                hs, control_rows,
                drive_client=drive_client,
                sop_folder_id=cfg.drive.sop_folder_id,
            )

            overwrite = overwrite_existing_override or cfg.app.overwrite_existing
            modified_bytes, write_result = write_to_template(
                doc_bytes=current_bytes,
                table_match_info={"table_index": -1, "invisible_code": mrow.invisible_code},
                control_row=control_row,
                overwrite_existing=overwrite,
                update_only_blank_cells=cfg.app.update_only_blank_cells,
                allow_price_update=cfg.app.allow_price_update_when_filling_blanks,
                strict_mode=cfg.app.strict_mode,
                remove_invisible_code=False,  # Don't remove yet - more rows to write
                header_row_1based=mrow.header_row,
            )

            if modified_bytes is None:
                msg = f"DOCX write failed for HS #{hs.homesite}: {write_result.error}"
                logger.error(msg)
                errors_per_hs.append({"homesite": hs.homesite, "plan": hs.plan, "error": msg})
                error_count += 1
                continue

            current_bytes = modified_bytes
            write_actions.append({
                "homesite": hs.homesite,
                "plan": hs.plan,
                "action": write_result.action,
                "row": write_result.row_index + 1,
                "price": control_row.price,
            })
            success_count += 1

        if not write_actions:
            # All homesites failed for this template
            continue

        # Now remove invisible code from the final version
        if cfg.app.remove_invisible_code:
            from .locator import find_table_by_invisible_code as find_ic, remove_invisible_code as do_remove
            import io as _io
            from docx import Document
            try:
                doc = Document(_io.BytesIO(current_bytes))
                match = find_ic(doc, mrow.invisible_code)
                do_remove(match, mrow.invisible_code)
                output = _io.BytesIO()
                doc.save(output)
                current_bytes = output.getvalue()
            except Exception as e:
                logger.warning("Could not remove invisible code from '%s': %s", template_name, e)

        # Check for changes (hash comparison)
        new_hash = compute_hash(current_bytes)
        old_hash_key = f"{pdf_id}_{template_name}"
        old_hash = manifest.get(old_hash_key, {}).get("output_hash", "")
        if cfg.app.skip_unchanged and new_hash == old_hash and old_hash:
            logger.info("Output unchanged for template '%s', skipping upload.", template_name)
            continue

        if cfg.app.dry_run:
            for wa in write_actions:
                logger.info(
                    "[DRY RUN] Would %s HS #%s at row %d in '%s'",
                    wa["action"], wa["homesite"], wa["row"], template_name,
                )
            result["homesite_results"].extend(write_actions)
            continue

        # Generate PDF from the populated DOCX
        base_name = os.path.splitext(template_name)[0]
        docx_output_name = f"{base_name}.docx"
        pdf_output_name = f"{base_name}.pdf"

        try:
            output_pdf_bytes = export_to_pdf(
                drive_client, current_bytes,
                cfg.drive.final_price_sheets_folder_id,
                temp_name=f"_temp_{base_name}",
            )
        except Exception as e:
            logger.error("PDF export failed for '%s': %s", template_name, e)
            for hs in hs_list:
                errors_per_hs.append({"homesite": hs.homesite, "plan": hs.plan, "error": f"PDF export: {e}"})
            error_count += len(hs_list)
            continue

        # Upload DOCX + PDF with safe replace
        final_folder = cfg.drive.final_price_sheets_folder_id
        try:
            docx_result = drive_client.safe_replace(
                current_bytes, final_folder, docx_output_name,
                allow_deletions=cfg.drive.allow_deletions,
            )
            pdf_result_file = drive_client.safe_replace(
                output_pdf_bytes, final_folder, pdf_output_name,
                allow_deletions=cfg.drive.allow_deletions,
            )

            all_output_ids[template_name] = {
                "docx_id": docx_result["id"],
                "pdf_id": pdf_result_file["id"],
            }
        except Exception as e:
            logger.error("Upload failed for '%s': %s", template_name, e)
            error_count += len(hs_list)
            continue

        # Record in manifest per template
        manifest[old_hash_key] = {
            "name": pdf_name,
            "template": template_name,
            "processed_at": datetime.now(timezone.utc).isoformat(),
            "output_hash": new_hash,
            "output_ids": all_output_ids.get(template_name, {}),
            "community": community,
            "phase": phase,
            "homesites": [wa["homesite"] for wa in write_actions],
            "actions": write_actions,
        }

        result["homesite_results"].extend(write_actions)

        for wa in write_actions:
            log_event(
                logger, "INFO", "Homesite written to template",
                pdf_name=pdf_name, template=template_name,
                homesite=wa["homesite"], plan=wa["plan"],
                action=wa["action"], row=wa["row"],
            )

    # Step 6: Mark the PDF as processed
    if success_count > 0 and not cfg.app.dry_run:
        try:
            drive_client.set_app_properties(pdf_id, {
                "processed": "true",
                "processed_at": datetime.now(timezone.utc).isoformat(),
                "processed_by": platform.node(),
                "community": community,
                "phase": phase,
                "homesite_count": str(success_count),
            })
        except Exception as e:
            logger.warning("Failed to set appProperties on PDF: %s", e)

        manifest[pdf_id] = {
            "name": pdf_name,
            "processed_at": datetime.now(timezone.utc).isoformat(),
            "community": community,
            "phase": phase,
            "homesite_count": success_count,
            "error_count": error_count,
        }

    # Move/copy processed PDF
    if success_count > 0 and cfg.drive.move_processed_pdfs and not cfg.app.dry_run:
        try:
            if not cfg.drive.keep_originals:
                drive_client.move_file(
                    pdf_id, cfg.drive.final_price_sheets_folder_id,
                    cfg.drive.new_releases_folder_id,
                )
        except Exception as e:
            logger.warning("Failed to move processed PDF: %s", e)

    result["output_ids"] = all_output_ids
    result["homesite_results"].extend(errors_per_hs)

    if error_count == 0 and success_count > 0:
        result["status"] = "success"
        result["details"] = f"Wrote {success_count} homesite(s) across {len(template_groups)} template(s)."
    elif success_count > 0:
        result["status"] = "partial"
        result["details"] = f"{success_count} OK, {error_count} errors."
    elif cfg.app.dry_run:
        result["status"] = "dry_run"
        result["details"] = f"Would write {len(write_actions)} homesite(s) (dry run)."
    else:
        result["status"] = "error"
        result["details"] = f"All {error_count} homesite(s) failed."

    log_event(
        logger, "INFO", "Release PDF processing complete",
        pdf_name=pdf_name, status=result["status"],
        success=success_count, errors=error_count,
    )
    return result


# ── Legacy single-homesite processing (kept for backward compatibility) ──

def process_single_pdf(
    cfg: Config,
    drive_client: DriveClient,
    sheets_client: SheetsClient,
    control_rows: list,
    mapping_rows: list,
    pdf_file: dict,
    manifest: dict,
    certs: dict,
    overwrite_existing_override: bool = False,
) -> dict:
    """Process a single PDF using the OLD filename-based flow.

    This is the legacy path for PDFs named like Community_Homesite_Floorplan.pdf.
    For release PDFs (Community Phase XX.pdf), use process_release_pdf instead.
    """
    pdf_name = pdf_file["name"]
    pdf_id = pdf_file["id"]
    result = {
        "pdf_name": pdf_name, "pdf_id": pdf_id,
        "status": "unknown", "details": "", "output_ids": {},
    }

    log_event(logger, "INFO", "Processing PDF (legacy)", pdf_name=pdf_name, pdf_id=pdf_id)

    # Parse filename for community/homesite/floorplan
    parsed = parse_pdf_filename(pdf_name)
    if not parsed:
        result["status"] = "error"
        result["details"] = f"Cannot parse community/homesite/floorplan from filename: {pdf_name}"
        quarantine_pdf(drive_client, pdf_file, cfg.drive.final_price_sheets_folder_id,
                       "unparseable_filename", cfg.pdf.quarantine_folder_name)
        return result

    community = parsed["community"]
    homesite = parsed["homesite"]
    floorplan = parsed["floorplan"]
    result["community"] = community
    result["homesite"] = homesite
    result["floorplan"] = floorplan

    logger.info("Parsed: community=%s homesite=%s floorplan=%s", community, homesite, floorplan)

    # Find CONTROL row
    control_row = find_control_row(control_rows, community, homesite, floorplan)
    if not control_row:
        if cfg.pdf.require_control_match:
            result["status"] = "error"
            result["details"] = f"No matching CONTROL row for ({community}, {homesite}, {floorplan})"
            quarantine_pdf(drive_client, pdf_file, cfg.drive.final_price_sheets_folder_id,
                           "no_control_match", cfg.pdf.quarantine_folder_name)
            return result
        else:
            result["status"] = "skipped"
            result["details"] = "No CONTROL match (not required)"
            return result

    # SOP address resolution
    if not control_row.address.strip():
        addr = resolve_address(drive_client, cfg.drive.sop_folder_id, community, homesite, floorplan)
        if addr:
            control_row.address = addr

    # Find MAPPING row
    mapping_row = find_mapping_row(mapping_rows, community, floorplan)
    if not mapping_row:
        result["status"] = "error"
        result["details"] = f"No MAPPING row for ({community}, {floorplan})"
        quarantine_pdf(drive_client, pdf_file, cfg.drive.final_price_sheets_folder_id,
                       "no_mapping_match", cfg.pdf.quarantine_folder_name)
        return result

    template_name = mapping_row.file_name
    invisible_code = mapping_row.invisible_code

    # Download template
    template_file = drive_client.find_file_by_name(cfg.drive.templates_folder_id, template_name)
    if not template_file:
        result["status"] = "error"
        result["details"] = f"Template '{template_name}' not found in Templates folder."
        quarantine_pdf(drive_client, pdf_file, cfg.drive.final_price_sheets_folder_id,
                       "template_not_found", cfg.pdf.quarantine_folder_name)
        return result

    if not cfg.app.allow_uncertified_templates:
        if not is_template_certified(template_file["id"], template_file.get("modifiedTime", ""), certs):
            result["status"] = "error"
            result["details"] = (
                f"Template '{template_name}' is not certified. "
                f"Run --certify-template --community {community} --floorplan {floorplan}"
            )
            return result

    cache_path = os.path.join(cfg.drive.download_cache_dir, template_name)
    drive_client.download_file(template_file["id"], cache_path)
    with open(cache_path, "rb") as f:
        template_bytes = f.read()

    overwrite = overwrite_existing_override or cfg.app.overwrite_existing
    modified_bytes, write_result = write_to_template(
        doc_bytes=template_bytes,
        table_match_info={"table_index": -1, "invisible_code": invisible_code},
        control_row=control_row,
        overwrite_existing=overwrite,
        update_only_blank_cells=cfg.app.update_only_blank_cells,
        allow_price_update=cfg.app.allow_price_update_when_filling_blanks,
        strict_mode=cfg.app.strict_mode,
        remove_invisible_code=cfg.app.remove_invisible_code,
        header_row_1based=mapping_row.header_row,
    )

    if modified_bytes is None:
        result["status"] = "error"
        result["details"] = f"DOCX write failed: {write_result.error}"
        reason = write_result.action if write_result.action else "docx_write_failed"
        quarantine_pdf(drive_client, pdf_file, cfg.drive.final_price_sheets_folder_id,
                       reason, cfg.pdf.quarantine_folder_name)
        return result

    new_hash = compute_hash(modified_bytes)
    old_hash = manifest.get(pdf_id, {}).get("output_hash", "")
    if cfg.app.skip_unchanged and new_hash == old_hash and old_hash:
        result["status"] = "skipped_unchanged"
        result["details"] = "Output unchanged, skipping upload."
        return result

    if cfg.app.dry_run:
        result["status"] = "dry_run"
        result["details"] = f"Would {write_result.action} at row {write_result.row_index + 1}"
        return result

    base_name = os.path.splitext(template_name)[0]
    docx_output_name = f"{base_name}.docx"
    pdf_output_name = f"{base_name}.pdf"

    pdf_bytes = export_to_pdf(
        drive_client, modified_bytes,
        cfg.drive.final_price_sheets_folder_id,
        temp_name=f"_temp_{base_name}",
    )

    final_folder = cfg.drive.final_price_sheets_folder_id
    docx_result = drive_client.safe_replace(
        modified_bytes, final_folder, docx_output_name,
        allow_deletions=cfg.drive.allow_deletions,
    )
    pdf_result_file = drive_client.safe_replace(
        pdf_bytes, final_folder, pdf_output_name,
        allow_deletions=cfg.drive.allow_deletions,
    )

    result["output_ids"] = {
        "docx_id": docx_result["id"],
        "pdf_id": pdf_result_file["id"],
    }

    drive_client.set_app_properties(pdf_id, {
        "processed": "true",
        "processed_at": datetime.now(timezone.utc).isoformat(),
        "processed_by": platform.node(),
    })

    manifest[pdf_id] = {
        "name": pdf_name,
        "processed_at": datetime.now(timezone.utc).isoformat(),
        "output_hash": new_hash,
        "output_ids": result["output_ids"],
        "community": community,
        "homesite": homesite,
        "floorplan": floorplan,
        "result": write_result.action,
    }

    if cfg.drive.move_processed_pdfs:
        try:
            if not cfg.drive.keep_originals:
                drive_client.move_file(pdf_id, final_folder, cfg.drive.new_releases_folder_id)
        except Exception as e:
            logger.warning("Failed to move processed PDF: %s", e)

    result["status"] = "success"
    result["details"] = f"{write_result.action} at row {write_result.row_index + 1}"
    log_event(logger, "INFO", "PDF processed successfully",
              pdf_name=pdf_name, action=write_result.action, row=write_result.row_index + 1)
    return result


# ── Detect PDF type ──

def _is_release_pdf(filename: str) -> bool:
    """Detect whether a PDF is a multi-homesite release PDF.

    Release PDFs are named like 'Nova Phase 2D.pdf'.
    Legacy PDFs are named like 'Isla_101_2.pdf'.
    """
    parsed = parse_release_filename(filename)
    if parsed and parsed.get("phase"):
        return True  # Has a phase component → release format
    return False


# ── Run Process New Releases ──

def run_process_new_releases(
    cfg: Config,
    community_filter: str = None,
    homesite_filter: str = None,
    floorplan_filter: str = None,
    overwrite_existing_override: bool = False,
    once: bool = False,
):
    """Main entry point: process all new release PDFs."""
    if not acquire_lock():
        return

    try:
        # Connect
        sheets = SheetsClient(cfg.google.credentials_json_path, cfg.google.spreadsheet_id)
        sheets.connect()
        drive = DriveClient(cfg.google.credentials_json_path, cfg.drive.shared_drive_id)
        drive.connect()
        drive.connect_for_writes()  # OAuth2 for uploads

        cfg.ensure_cache_dirs()

        while True:
            # Load sheet data
            control_records = sheets.get_all_records(cfg.google.control_tab)
            control_rows = parse_control_tab(control_records)
            mapping_records = sheets.get_all_records(cfg.google.mapping_tab)
            mapping_rows = parse_mapping_tab(mapping_records)

            # Load manifest and certs
            manifest = load_manifest(cfg.drive.processed_manifest)
            certs = load_certifications()

            # List new PDFs
            pdfs = drive.list_pdfs(cfg.drive.new_releases_folder_id)
            logger.info("Found %d PDFs in New Releases.", len(pdfs))

            processed_count = 0
            error_count = 0

            for pdf_file in pdfs:
                # Skip already processed
                if is_already_processed(pdf_file, manifest):
                    logger.debug("Skipping already-processed: %s", pdf_file["name"])
                    continue

                # Determine PDF type and route accordingly
                if _is_release_pdf(pdf_file["name"]):
                    # ── New release PDF format (multi-homesite) ──
                    # Apply community filter from filename
                    if community_filter:
                        parsed_fn = parse_release_filename(pdf_file["name"])
                        if parsed_fn:
                            if normalize_for_compare(parsed_fn["community"]) != normalize_for_compare(community_filter):
                                continue

                    result = process_release_pdf(
                        cfg, drive, sheets, control_rows, mapping_rows,
                        pdf_file, manifest, certs,
                        overwrite_existing_override=overwrite_existing_override,
                        homesite_filter=homesite_filter,
                        floorplan_filter=floorplan_filter,
                    )
                else:
                    # ── Legacy filename-based format ──
                    parsed = parse_pdf_filename(pdf_file["name"])
                    if parsed and community_filter:
                        if normalize_for_compare(parsed["community"]) != normalize_for_compare(community_filter):
                            continue
                    if parsed and homesite_filter:
                        if normalize_for_compare(parsed["homesite"]) != normalize_for_compare(homesite_filter):
                            continue
                    if parsed and floorplan_filter:
                        if normalize_for_compare(parsed["floorplan"]) != normalize_for_compare(floorplan_filter):
                            continue

                    result = process_single_pdf(
                        cfg, drive, sheets, control_rows, mapping_rows,
                        pdf_file, manifest, certs, overwrite_existing_override,
                    )

                if result["status"] in ("success", "partial"):
                    processed_count += 1
                elif result["status"] == "error":
                    error_count += 1
                    logger.error("Error processing %s: %s", pdf_file["name"], result["details"])

                # Save manifest after each PDF
                save_manifest(cfg.drive.processed_manifest, manifest)

            logger.info(
                "Cycle complete: %d processed, %d errors, %d total PDFs.",
                processed_count, error_count, len(pdfs),
            )

            # Print summary
            if processed_count > 0 or error_count > 0:
                print(f"\nProcessed: {processed_count}  Errors: {error_count}  Total PDFs: {len(pdfs)}")

            # Polling or exit
            if cfg.app.poll_interval_seconds > 0 and not once:
                if cfg.app.print_watching_message:
                    print("Watching for new changes...")
                    logger.info("Watching for new changes...")
                time.sleep(cfg.app.poll_interval_seconds)
            else:
                break

    finally:
        release_lock()


# ── Sync CONTROL sheet -> templates ──

def sync_control_to_templates(
    cfg: Config,
    sheets_client: SheetsClient,
    drive_client: DriveClient,
    control_rows: list,
    mapping_rows: list,
    certs: dict,
) -> dict:
    """Push CONTROL sheet data into Word templates and re-upload.

    Each template file (e.g. Sella.docx) contains MULTIPLE tables, one per
    floorplan.  Each floorplan's table is identified by its own invisible
    code.  We group control rows by (template_file, mapping_row) so each
    floorplan's rows go to the correct table.

    Also resolves missing addresses from SOP for any CONTROL row that has
    a blank address field before writing to the template.

    Returns dict with counts: {templates_updated, rows_synced, errors}.
    """
    from .template_reader import read_final_docx_data

    result = {"templates_updated": 0, "rows_synced": 0, "errors": 0}

    if not control_rows:
        return result

    # ── Group control rows by (template_file, invisible_code) ──
    # Each (template_file, invisible_code) maps to one specific table inside
    # the DOCX.  We need to write each group to its own table.
    #
    # Structure: floorplan_groups[invisible_code] = {mapping_row, crows: []}
    # And:       file_groups[template_file_name] = {floorplan_groups, ...}

    file_groups = {}  # template_file_name -> {ic -> {mrow, crows}}

    # First, seed file_groups from ALL mapping rows (even those with no CONTROL
    # rows).  This ensures that if ALL rows for a template/floorplan are deleted
    # from CONTROL, we still process that template and clear the deleted rows.
    for mrow in mapping_rows:
        fname = mrow.file_name
        ic = mrow.invisible_code
        if fname not in file_groups:
            file_groups[fname] = {}
        if ic not in file_groups[fname]:
            file_groups[fname][ic] = {"mapping_row": mrow, "control_rows": []}

    # Then, add CONTROL rows to the appropriate groups
    for crow in control_rows:
        mrow = find_mapping_row(mapping_rows, crow.community, crow.floorplan)
        if not mrow:
            continue

        fname = mrow.file_name
        ic = mrow.invisible_code

        # Group should already exist from the mapping seed above
        if fname not in file_groups:
            file_groups[fname] = {}
        if ic not in file_groups[fname]:
            file_groups[fname][ic] = {"mapping_row": mrow, "control_rows": []}
        file_groups[fname][ic]["control_rows"].append(crow)

    if not file_groups:
        return result

    cfg.ensure_cache_dirs()

    # SOP address resolution cache to avoid re-resolving the same homesite
    _sop_cache = {}  # (community, homesite) -> address or ""

    def _resolve_sop(community, homesite, floorplan):
        key = (normalize_for_compare(community), normalize_for_compare(homesite))
        if key not in _sop_cache:
            try:
                addr = resolve_address(
                    drive_client, cfg.drive.sop_folder_id,
                    community, homesite, floorplan,
                )
                _sop_cache[key] = addr or ""
            except Exception:
                _sop_cache[key] = ""
        return _sop_cache[key]

    for template_name, fp_groups in file_groups.items():
        # Get any mapping row to check certification (same file)
        any_mrow = next(iter(fp_groups.values()))["mapping_row"]

        # Find the ORIGINAL blank template (has invisible codes)
        original_template = drive_client.find_file_by_name(
            cfg.drive.templates_folder_id, template_name
        )
        if not original_template:
            logger.warning("Template '%s' not found for sheet->template sync.", template_name)
            result["errors"] += 1
            continue

        if not cfg.app.allow_uncertified_templates:
            if not is_template_certified(
                original_template["id"], original_template.get("modifiedTime", ""), certs
            ):
                continue

        # ── Step A: Fill missing addresses from SOP for all CONTROL rows ──
        # Do this BEFORE the change-detection step so we know what will actually
        # be written (a row with a newly-resolved address counts as a change).
        for ic, fp_data in fp_groups.items():
            for crow in fp_data["control_rows"]:
                if not crow.address.strip():
                    sop_addr = _resolve_sop(crow.community, crow.homesite, crow.floorplan)
                    if sop_addr:
                        crow.address = sop_addr
                        print(f"    [SOP] HS#{crow.homesite} ({crow.community}): {sop_addr}")
                        logger.info(
                            "SOP resolved address for HS#%s (%s): %s",
                            crow.homesite, crow.community, sop_addr,
                        )

        # ── Step B: Read the FINAL price sheet to see what's currently in it ──
        base_name = os.path.splitext(template_name)[0]
        final_docx_name = f"{base_name}.docx"
        final_file = drive_client.find_file_by_name(
            cfg.drive.final_price_sheets_folder_id, final_docx_name
        )

        # Read ALL tables from final file for comparison
        existing_by_hs = {}  # homesite_norm -> row_data dict
        if final_file:
            final_cache_path = os.path.join(cfg.drive.download_cache_dir, f"_final_{template_name}")
            try:
                drive_client.download_file(final_file["id"], final_cache_path)
                with open(final_cache_path, "rb") as f:
                    final_bytes = f.read()

                # Read ALL floorplan tables from the final file
                for ic, fp_data in fp_groups.items():
                    mrow = fp_data["mapping_row"]
                    existing_rows = read_final_docx_data(
                        final_bytes,
                        mrow.community, mrow.floorplan,
                        header_row_1based=mrow.header_row,
                    )
                    for erow in existing_rows:
                        hs_key = normalize_for_compare(erow["homesite"])
                        existing_by_hs[hs_key] = erow

                logger.info(
                    "Read %d existing rows from final '%s' for comparison.",
                    len(existing_by_hs), final_docx_name,
                )
            except Exception as e:
                logger.info(
                    "Could not read final file '%s' for comparison: %s. Will rebuild.",
                    final_docx_name, e,
                )

        # ── Step C: Check if ANY control row has changes vs the final file ──
        any_changes = False
        deletion_detected = False

        # Build a set of all homesite keys currently in CONTROL for this template
        control_hs_keys = set()
        for ic, fp_data in fp_groups.items():
            for crow in fp_data["control_rows"]:
                control_hs_keys.add(normalize_for_compare(crow.homesite))

        # Check 1: Are there homesites in the FINAL file that are NOT in CONTROL?
        # (i.e., the user deleted a row from CONTROL — it must be removed from template)
        for hs_key in existing_by_hs:
            if hs_key not in control_hs_keys:
                any_changes = True
                deletion_detected = True
                logger.info(
                    "Row deletion detected: HS '%s' exists in final '%s' "
                    "but not in CONTROL sheet. Will rebuild.",
                    hs_key, template_name,
                )
                break

        # Check 2: Are there new or modified CONTROL rows vs what's in the final file?
        if not any_changes:
            for ic, fp_data in fp_groups.items():
                for crow in fp_data["control_rows"]:
                    hs_key = normalize_for_compare(crow.homesite)
                    erow = existing_by_hs.get(hs_key)

                    if erow is None:
                        # Homesite not yet in the final file at all
                        any_changes = True
                        break

                    # Notes: always push from CONTROL (this is the "sold" case)
                    if crow.notes.strip() != erow.get("notes", "").strip():
                        any_changes = True
                        break
                    # Price: push if CONTROL has a value and it differs
                    if crow.price.strip() and crow.price.strip() != erow.get("price", "").strip():
                        any_changes = True
                        break
                    # Address: push if CONTROL has an address (from sheet or SOP)
                    # that is different from what's in the template
                    if crow.address.strip() and crow.address.strip() != erow.get("address", "").strip():
                        any_changes = True
                        break
                    # Ready by: push if CONTROL has a value and it differs
                    if crow.ready_by.strip() and crow.ready_by.strip() != erow.get("ready_by", "").strip():
                        any_changes = True
                        break

                if any_changes:
                    break

        if not any_changes:
            logger.info("Template '%s' already in sync with CONTROL sheet.", template_name)
            continue  # This template is fully in sync

        # ── Step D: Rebuild from BLANK template ──
        # Download the blank template (which has ALL invisible codes intact)
        blank_cache_path = os.path.join(cfg.drive.download_cache_dir, f"_blank_{template_name}")
        try:
            drive_client.download_file(original_template["id"], blank_cache_path)
        except Exception as e:
            logger.warning("Failed to download blank template '%s': %s", template_name, e)
            result["errors"] += 1
            continue

        with open(blank_cache_path, "rb") as f:
            current_bytes = f.read()

        total_rows = sum(len(fp["control_rows"]) for fp in fp_groups.values())
        logger.info(
            "Sheet->template sync: rebuilding '%s' from blank template "
            "(%d floorplan table(s), %d total row(s))",
            template_name, len(fp_groups), total_rows,
        )

        # Write each floorplan group to its OWN table (via its own invisible code)
        rows_written = 0
        for ic, fp_data in fp_groups.items():
            mrow = fp_data["mapping_row"]
            crow_list = fp_data["control_rows"]

            for crow in crow_list:
                modified_bytes, write_result = write_to_template(
                    doc_bytes=current_bytes,
                    table_match_info={"table_index": -1, "invisible_code": ic},
                    control_row=crow,
                    overwrite_existing=True,
                    update_only_blank_cells=False,
                    allow_price_update=True,
                    strict_mode=False,
                    remove_invisible_code=False,  # Keep codes for next floorplan
                    header_row_1based=mrow.header_row,
                )

                if modified_bytes is None:
                    logger.warning(
                        "Sheet->template sync: write failed for HS#%s (%s/%s) in '%s': %s",
                        crow.homesite, crow.community, crow.floorplan,
                        template_name, write_result.error,
                    )
                    result["errors"] += 1
                    continue

                current_bytes = modified_bytes
                rows_written += 1

        if rows_written == 0 and not deletion_detected:
            continue

        if cfg.app.dry_run:
            if deletion_detected:
                print(f"  [DRY RUN] Would update '{template_name}' "
                      f"({rows_written} row(s), removed deleted rows)")
            else:
                print(f"  [DRY RUN] Would update '{template_name}' ({rows_written} row(s))")
            continue

        # Generate PDF and upload both
        try:
            output_pdf_bytes = export_to_pdf(
                drive_client, current_bytes,
                cfg.drive.final_price_sheets_folder_id,
                temp_name=f"_sync_{base_name}",
            )

            final_folder = cfg.drive.final_price_sheets_folder_id
            drive_client.safe_replace(
                current_bytes, final_folder, f"{base_name}.docx",
                allow_deletions=cfg.drive.allow_deletions,
            )
            drive_client.safe_replace(
                output_pdf_bytes, final_folder, f"{base_name}.pdf",
                allow_deletions=cfg.drive.allow_deletions,
            )

            result["templates_updated"] += 1
            result["rows_synced"] += rows_written
            if deletion_detected:
                print(f"  Updated '{template_name}': {rows_written} row(s), "
                      f"removed deleted rows from template")
            else:
                print(f"  Updated '{template_name}': {rows_written} row(s) across "
                      f"{len(fp_groups)} floorplan table(s)")

        except Exception as e:
            logger.error("Failed to upload synced template '%s': %s", template_name, e)
            result["errors"] += 1

    return result


# ══════════════════════════════════════════════════════════════
#  MASTER COMMAND  --  One command does EVERYTHING
# ══════════════════════════════════════════════════════════════

def run_master(
    cfg: Config,
    community_filter: str = None,
    homesite_filter: str = None,
    floorplan_filter: str = None,
    overwrite_existing_override: bool = False,
    once: bool = False,
):
    """Master command: runs the entire pipeline in one shot.

    The CONTROL spreadsheet is the single source of truth.
    Data flows: New PDFs -> CONTROL sheet -> Templates -> Final PDFs.
    Templates NEVER write back to CONTROL (to preserve user deletions/edits).

    Steps:
      1. Health check (verify connections)
      2. Scan New Releases folder for PDFs
      3. Parse each PDF -> extract homesites, SOP addresses -> update CONTROL
      4. Certify all templates
      5. Sync CONTROL sheet -> templates (push all data, handle deletions)
      6. Process new PDFs into templates + upload DOCX/PDF
      7. Summary
    """

    print("=" * 60)
    print("  PRICE SHEET BOT - MASTER RUN")
    print("=" * 60)

    # ── Step 1: Health Check ──
    print("\n[STEP 1/7] Health Check")
    print("-" * 40)
    ok = run_health_check(cfg)
    if not ok:
        print("\nHealth check FAILED. Fix issues above first.")
        return

    # ── Connect everything ──
    sheets = SheetsClient(cfg.google.credentials_json_path, cfg.google.spreadsheet_id)
    sheets.connect()
    drive = DriveClient(cfg.google.credentials_json_path, cfg.drive.shared_drive_id)
    drive.connect()
    drive.connect_for_writes()
    cfg.ensure_cache_dirs()

    mapping_records = sheets.get_all_records(cfg.google.mapping_tab)
    mapping_rows = parse_mapping_tab(mapping_records)

    # ── Step 2: Scan New Releases ──
    print(f"\n[STEP 2/7] Scanning New Releases folder")
    print("-" * 40)
    pdfs = drive.list_pdfs(cfg.drive.new_releases_folder_id)
    manifest = load_manifest(cfg.drive.processed_manifest)

    # Filter out already-processed
    new_pdfs = [p for p in pdfs if not is_already_processed(p, manifest)]

    # Apply community filter from filename
    if community_filter:
        filtered = []
        for p in new_pdfs:
            rel = parse_release_filename(p["name"])
            if rel and normalize_for_compare(rel.get("community", "")) == normalize_for_compare(community_filter):
                filtered.append(p)
            else:
                legacy = parse_pdf_filename(p["name"])
                if legacy and normalize_for_compare(legacy.get("community", "")) == normalize_for_compare(community_filter):
                    filtered.append(p)
        new_pdfs = filtered

    print(f"  Total PDFs in folder: {len(pdfs)}")
    print(f"  New (unprocessed): {len(new_pdfs)}")
    for p in new_pdfs:
        print(f"    - {p['name']}")

    # ── Step 3: Parse PDFs + SOP addresses + Update CONTROL tab ──
    # New PDFs feed INTO the CONTROL sheet. The sheet is the authority.
    print(f"\n[STEP 3/7] Parsing PDFs, resolving addresses, updating Google Sheet")
    print("-" * 40)

    all_homesites = []  # List of (pdf_file, ReleaseHomesite) tuples
    pdf_parse_results = {}  # pdf_id -> ParsedReleasePDF
    pdf_sheet_count = 0
    pdf_upsert_rows = []  # Collect rows for batch upsert

    if not new_pdfs:
        print("  No new PDFs to parse.")
    else:
        for pdf_file in new_pdfs:
            pdf_name = pdf_file["name"]

            if _is_release_pdf(pdf_name):
                # Download and parse the release PDF
                local_path = os.path.join(cfg.drive.download_cache_dir, pdf_name)
                try:
                    drive.download_file(pdf_file["id"], local_path)
                except Exception as e:
                    print(f"  ERROR downloading '{pdf_name}': {e}")
                    continue

                parsed = parse_release_pdf(local_path)
                pdf_parse_results[pdf_file["id"]] = parsed

                if not parsed.homesites:
                    print(f"  WARNING: No homesites found in '{pdf_name}'")
                    continue

                community = parsed.meta.community
                phase = parsed.meta.phase
                print(f"\n  {pdf_name}: community={community}, phase={phase}, {len(parsed.homesites)} homesites")

                for hs in parsed.homesites:
                    # Apply homesite/floorplan filters
                    if homesite_filter and normalize_for_compare(hs.homesite) != normalize_for_compare(homesite_filter):
                        continue
                    if floorplan_filter and normalize_for_compare(hs.plan) != normalize_for_compare(floorplan_filter):
                        continue

                    # Resolve address from SOP
                    address = ""
                    try:
                        addr = resolve_address(drive, cfg.drive.sop_folder_id, community, hs.homesite, hs.plan)
                        if addr:
                            address = addr
                    except Exception:
                        pass

                    # Determine price (Total Released Price)
                    price = hs.total_released_price if hs.total_released_price else hs.base_price

                    # COE / ready_by: raw date from PDF (e.g. "April, 2026")
                    # Both ready_by and move_in get the same raw value here;
                    # parse_ready_by() will convert it when writing to the DOCX.
                    ready_by = hs.coe_date if hs.coe_date else hs.default_coe

                    # Collect for batch upsert
                    pdf_upsert_rows.append({
                        "community": community,
                        "homesite": hs.homesite,
                        "floorplan": hs.plan,
                        "price": price,
                        "address": address,
                        "ready_by": ready_by,
                        "move_in": ready_by,  # human-readable date shown in sheet
                        "notes": "",
                    })

                    print(f"    HS #{hs.homesite} Plan {hs.plan}: {price}"
                          f"  COE: {ready_by}"
                          f"  Addr: {address or '(none)'}")

                    all_homesites.append((pdf_file, hs))
            else:
                print(f"  {pdf_name}: legacy format (will process directly)")

    # Batch upsert all PDF-extracted rows to CONTROL tab
    if pdf_upsert_rows:
        print(f"\n  Writing {len(pdf_upsert_rows)} row(s) to CONTROL tab (batch mode)...")
        try:
            pdf_results = sheets.batch_upsert_control_rows(
                cfg.google.control_tab, pdf_upsert_rows
            )
            for urow, action in pdf_results:
                symbol = {"inserted": "+", "updated": "~", "skipped": "="}
                print(f"    [{symbol.get(action, '?')}] {urow['community']} HS#{urow['homesite']} "
                      f"Plan {urow['floorplan']} -> {action}")
                pdf_sheet_count += 1
        except Exception as e:
            print(f"    [!] Batch sheet update failed: {e}")

    print(f"\n  CONTROL tab: {pdf_sheet_count} row(s) from PDFs.")

    # ── Step 3b: Resolve SOP addresses for existing CONTROL rows with blank address ──
    # This runs regardless of whether there were new PDFs, so rows already in the
    # sheet that still have no address get filled in from SOP.
    print(f"\n  Resolving missing addresses from SOP for existing CONTROL rows...")
    try:
        existing_control_records = sheets.get_all_records(cfg.google.control_tab)
        existing_control_rows = parse_control_tab(existing_control_records)
        sop_address_updates = []
        for crow in existing_control_rows:
            if not crow.address.strip():
                try:
                    addr = resolve_address(
                        drive, cfg.drive.sop_folder_id,
                        crow.community, crow.homesite, crow.floorplan,
                    )
                    if addr:
                        sop_address_updates.append({
                            "community": crow.community,
                            "homesite": crow.homesite,
                            "floorplan": crow.floorplan,
                            "price": crow.price,
                            "address": addr,
                            "ready_by": crow.ready_by,
                            "move_in": crow.move_in,
                            "notes": crow.notes,
                        })
                        print(f"    [SOP] {crow.community} HS#{crow.homesite}: {addr}")
                except Exception:
                    pass

        if sop_address_updates:
            sheets.batch_upsert_control_rows(cfg.google.control_tab, sop_address_updates)
            print(f"  Updated {len(sop_address_updates)} address(es) from SOP.")
        else:
            print("  No missing addresses found (or SOP has no matches).")
    except Exception as e:
        print(f"  WARNING: SOP address resolution failed: {e}")

    # ── Step 4: Certify all templates ──
    print(f"\n[STEP 4/7] Certifying templates")
    print("-" * 40)
    cert_ok, cert_total, cert_pass, cert_fail = run_certify_all(cfg, sheets=sheets, drive=drive)
    if cert_fail > 0:
        print(f"  WARNING: {cert_fail} template(s) failed certification (will be skipped).")

    # ── Step 5: Sync CONTROL sheet -> templates ──
    # The CONTROL sheet is the single source of truth.
    # - Rows in CONTROL are written to templates
    # - Rows deleted from CONTROL are removed from templates
    # - Formatting (sold, upgraded flooring) is applied based on notes
    print(f"\n[STEP 5/7] Syncing CONTROL sheet -> templates (CONTROL is authority)")
    print("-" * 40)

    # Reload CONTROL data (it was updated in step 3)
    control_records = sheets.get_all_records(cfg.google.control_tab)
    control_rows = parse_control_tab(control_records)
    mapping_records = sheets.get_all_records(cfg.google.mapping_tab)
    mapping_rows_fresh = parse_mapping_tab(mapping_records)
    certs = load_certifications()

    sheet_sync = sync_control_to_templates(
        cfg, sheets, drive, control_rows, mapping_rows_fresh, certs,
    )
    if sheet_sync["templates_updated"] > 0:
        print(f"  {sheet_sync['templates_updated']} template(s) updated, "
              f"{sheet_sync['rows_synced']} row(s) synced.")
    else:
        print("  All templates already in sync with CONTROL sheet.")
    if sheet_sync["errors"] > 0:
        print(f"  WARNING: {sheet_sync['errors']} error(s) during sync.")

    # ── Step 6: Process new PDFs -> write to templates + upload ──
    print(f"\n[STEP 6/7] Writing new PDFs to templates and generating output files")
    print("-" * 40)

    processed_count = 0
    error_count = 0

    if not new_pdfs:
        print("  No new PDFs to process.")
    else:
        if not acquire_lock():
            print("  ERROR: Could not acquire process lock.")
            return

        try:
            # Reuse control_rows and mapping_rows_fresh from Step 5
            manifest = load_manifest(cfg.drive.processed_manifest)

            for pdf_file in new_pdfs:
                if is_already_processed(pdf_file, manifest):
                    continue

                if _is_release_pdf(pdf_file["name"]):
                    result = process_release_pdf(
                        cfg, drive, sheets, control_rows, mapping_rows_fresh,
                        pdf_file, manifest, certs,
                        overwrite_existing_override=overwrite_existing_override,
                        homesite_filter=homesite_filter,
                        floorplan_filter=floorplan_filter,
                    )
                else:
                    result = process_single_pdf(
                        cfg, drive, sheets, control_rows, mapping_rows_fresh,
                        pdf_file, manifest, certs, overwrite_existing_override,
                    )

                status = result["status"]
                if status in ("success", "partial"):
                    processed_count += 1
                    print(f"  OK: {pdf_file['name']} -> {result['details']}")
                elif status == "error":
                    error_count += 1
                    print(f"  ERROR: {pdf_file['name']} -> {result['details']}")
                elif status == "dry_run":
                    print(f"  [DRY RUN]: {pdf_file['name']} -> {result['details']}")
                else:
                    print(f"  {status}: {pdf_file['name']} -> {result['details']}")

                save_manifest(cfg.drive.processed_manifest, manifest)

        finally:
            release_lock()

    # ── Step 7: Summary ──
    print(f"\n[STEP 7/7] Summary")
    print("-" * 40)
    print(f"  PDF -> Sheet sync:       {pdf_sheet_count} row(s)")
    print(f"  Templates certified:     {cert_pass}/{cert_total}")
    print(f"  Sheet -> Template sync:  {sheet_sync['templates_updated']} template(s), "
          f"{sheet_sync['rows_synced']} row(s)")
    print(f"  PDFs processed:          {processed_count}")
    print(f"  Errors:                  {error_count}")

    print("\n" + "=" * 60)
    print("  ALL DONE!")
    print("=" * 60)


# ── Health Check ──

def run_health_check(cfg: Config):
    """Verify all connections and permissions."""
    print("=== Price Sheet Bot Health Check ===\n")
    all_ok = True

    # 1. Sheet access
    print("[1] Checking Google Sheet access...")
    sheets = SheetsClient(cfg.google.credentials_json_path, cfg.google.spreadsheet_id)
    if sheets.verify_connection():
        print("    OK: Sheet accessible.")
    else:
        print("    FAIL: Cannot access sheet.")
        all_ok = False

    # 2. Drive access
    print("[2] Checking Google Drive access...")
    drive = DriveClient(cfg.google.credentials_json_path, cfg.drive.shared_drive_id)
    drive.connect()
    try:
        drive.connect_for_writes()
        print("    OK: OAuth2 user login active for uploads.")
    except Exception as e:
        print(f"    WARN: OAuth2 not configured ({e}). Uploads may fail on personal Drive.")

    folders = {
        "Templates": cfg.drive.templates_folder_id,
        "New Releases": cfg.drive.new_releases_folder_id,
        "Final Price Sheets": cfg.drive.final_price_sheets_folder_id,
        "SOP": cfg.drive.sop_folder_id,
    }
    for label, fid in folders.items():
        if drive.verify_folder_access(fid, label):
            print(f"    OK: {label} folder accessible.")
        else:
            print(f"    FAIL: {label} folder NOT accessible.")
            all_ok = False

    # 3. Upload ability
    print("[3] Checking upload ability to Final Price Sheets...")
    if drive.verify_upload_ability(cfg.drive.final_price_sheets_folder_id):
        print("    OK: Can upload to Final Price Sheets.")
    else:
        print("    FAIL: Cannot upload to Final Price Sheets.")
        all_ok = False

    # 4. Sheet tabs
    print("[4] Checking sheet tabs...")
    try:
        sheets.connect()
        for tab_name in [cfg.google.control_tab, cfg.google.mapping_tab]:
            records = sheets.get_all_records(tab_name)
            print(f"    OK: Tab '{tab_name}' has {len(records)} data rows.")
    except Exception as e:
        print(f"    FAIL: {e}")
        all_ok = False

    # 5. Templates exist
    print("[5] Checking template files...")
    try:
        mapping_records = sheets.get_all_records(cfg.google.mapping_tab)
        mapping_rows = parse_mapping_tab(mapping_records)
        for mrow in mapping_rows:
            tf = drive.find_file_by_name(cfg.drive.templates_folder_id, mrow.file_name)
            if tf:
                print(f"    OK: Template '{mrow.file_name}' found (id={tf['id']}).")
            else:
                print(f"    WARN: Template '{mrow.file_name}' NOT found in Drive.")
    except Exception as e:
        print(f"    FAIL: {e}")
        all_ok = False

    # 6. PDF parsing check
    print("[6] Checking PDF parser (pdfplumber)...")
    try:
        import pdfplumber
        print(f"    OK: pdfplumber {pdfplumber.__version__} available.")
    except ImportError:
        print("    FAIL: pdfplumber not installed. Run: pip install pdfplumber")
        all_ok = False

    print()
    if all_ok:
        print("Health check PASSED. All systems operational.")
    else:
        print("Health check FAILED. Fix issues above before processing.")
    return all_ok


# ── Audit Report ──

def run_audit_report(cfg: Config):
    """Print audit information."""
    print("=== Price Sheet Bot Audit Report ===\n")

    manifest = load_manifest(cfg.drive.processed_manifest)
    print(f"Processed PDFs in manifest: {len(manifest)}")
    for pid, info in manifest.items():
        if "homesite_count" in info:
            print(f"  - {info.get('name', pid)}: {info.get('homesite_count')} homesites at {info.get('processed_at', '?')}")
        else:
            print(f"  - {info.get('name', pid)}: {info.get('result', '?')} at {info.get('processed_at', '?')}")

    certs = load_certifications()
    print(f"\nCertified templates: {len(certs)}")
    for fid, info in certs.items():
        print(f"  - {info.get('file_name', fid)}: certified at {info.get('certified_at', '?')}")

    print()


# ── Template Certification ──

def run_certify_template(cfg: Config, community: str, floorplan: str,
                         sheets: SheetsClient = None, drive: DriveClient = None):
    """Certify a template for production use.

    Optionally pass existing sheets/drive connections to avoid rate limits.
    """
    print(f"Certifying template for community={community} floorplan={floorplan}...\n")

    if sheets is None:
        sheets = SheetsClient(cfg.google.credentials_json_path, cfg.google.spreadsheet_id)
        sheets.connect()
    if drive is None:
        drive = DriveClient(cfg.google.credentials_json_path, cfg.drive.shared_drive_id)
        drive.connect()
        drive.connect_for_writes()

    # Find mapping
    mapping_records = sheets.get_all_records(cfg.google.mapping_tab)
    mapping_rows = parse_mapping_tab(mapping_records)
    mrow = find_mapping_row(mapping_rows, community, floorplan)
    if not mrow:
        print(f"FAIL: No MAPPING row for ({community}, {floorplan}).")
        return False

    print(f"[1] Mapping found: file_name={mrow.file_name}, invisible_code={mrow.invisible_code}")

    # Find template
    tf = drive.find_file_by_name(cfg.drive.templates_folder_id, mrow.file_name)
    if not tf:
        print(f"FAIL: Template '{mrow.file_name}' not found in Drive.")
        return False

    print(f"[2] Template found in Drive: id={tf['id']}")

    # Download and verify
    cache_path = os.path.join(cfg.drive.download_cache_dir, mrow.file_name)
    cfg.ensure_cache_dirs()
    drive.download_file(tf["id"], cache_path)

    with open(cache_path, "rb") as f:
        template_bytes = f.read()

    from docx import Document
    import io as _io
    doc = Document(_io.BytesIO(template_bytes))

    # Check invisible code
    match = find_table_by_invisible_code(doc, mrow.invisible_code)
    print(f"[3] Invisible code found in table[{match.table_index}] at cell({match.cell_row},{match.cell_col})")

    # Check header row
    from .utils import build_header_map, validate_headers
    header_idx = mrow.header_row - 1
    if header_idx >= len(match.table.rows):
        print(f"FAIL: Header row {mrow.header_row} beyond table size.")
        return False

    header_cells = [cell.text.strip() for cell in match.table.rows[header_idx].cells]
    hmap = build_header_map(header_cells)
    print(f"[4] Header row {mrow.header_row}: {header_cells}")
    print(f"    Header map: {hmap}")

    missing = validate_headers(hmap)
    if missing:
        print(f"FAIL: Missing required headers: {missing}")
        return False
    print("[5] Required headers present.")

    # Check next blank row
    from .docx_writer import _find_next_blank_row
    data_start = header_idx + 1
    blank = _find_next_blank_row(match.table, hmap, data_start)
    if blank >= 0:
        print(f"[6] Next blank row: {blank + 1} (0-based: {blank})")
    else:
        print("[6] WARNING: No blank rows found in table.")

    # Test write/reopen
    try:
        output = _io.BytesIO()
        doc.save(output)
        reopened = Document(_io.BytesIO(output.getvalue()))
        print("[7] DOCX write/reopen: OK")
    except Exception as e:
        print(f"FAIL: DOCX write/reopen failed: {e}")
        return False

    # Test PDF export
    try:
        pdf_bytes = export_to_pdf(
            drive, output.getvalue(),
            cfg.drive.final_price_sheets_folder_id,
            temp_name=f"_cert_test_{mrow.file_name}",
        )
        print(f"[8] PDF export: OK ({len(pdf_bytes)} bytes)")
    except Exception as e:
        print(f"FAIL: PDF export failed: {e}")
        return False

    # Save certification
    certs = load_certifications()
    certs[tf["id"]] = {
        "file_name": mrow.file_name,
        "file_id": tf["id"],
        "modifiedTime": tf.get("modifiedTime", ""),
        "community": community,
        "floorplan": floorplan,
        "invisible_code": mrow.invisible_code,
        "certified_at": datetime.now(timezone.utc).isoformat(),
        "certified_by": platform.node(),
    }
    save_certifications(certs)

    print(f"\nTemplate '{mrow.file_name}' CERTIFIED for ({community}, {floorplan}).")
    return True


def run_certify_all(cfg: Config, sheets: SheetsClient = None,
                    drive: DriveClient = None) -> tuple:
    """Certify ALL templates listed in the MAPPING tab.

    Reads the MAPPING tab, loops through each unique (community, floorplan),
    and certifies each one. Already-certified templates that haven't changed
    are skipped quickly.

    Optionally pass existing sheets/drive connections to avoid rate limits.

    Returns: (all_ok: bool, total: int, passed: int, failed: int)
    """
    print("Certifying ALL templates from MAPPING tab...\n")

    if sheets is None:
        sheets = SheetsClient(cfg.google.credentials_json_path, cfg.google.spreadsheet_id)
        sheets.connect()
    if drive is None:
        drive = DriveClient(cfg.google.credentials_json_path, cfg.drive.shared_drive_id)
        drive.connect()
        drive.connect_for_writes()

    mapping_records = sheets.get_all_records(cfg.google.mapping_tab)
    mapping_rows = parse_mapping_tab(mapping_records)

    if not mapping_rows:
        print("No MAPPING rows found. Nothing to certify.")
        return (True, 0, 0, 0)

    # Deduplicate: one template may serve multiple floorplans, but we certify
    # per (community, floorplan) pair as that's how MAPPING is keyed.
    certs = load_certifications()
    total = len(mapping_rows)
    passed = 0
    failed = 0
    skipped = 0

    for i, mrow in enumerate(mapping_rows, 1):
        label = f"({mrow.community}, {mrow.floorplan})"
        print(f"--- [{i}/{total}] {label} -> {mrow.file_name} ---")

        # Quick check: is it already certified and unchanged?
        tf = drive.find_file_by_name(cfg.drive.templates_folder_id, mrow.file_name)
        if tf and is_template_certified(tf["id"], tf.get("modifiedTime", ""), certs):
            print(f"  Already certified. Skipping.\n")
            passed += 1
            skipped += 1
            continue

        # Run full certification (pass existing connections to avoid rate limits)
        try:
            ok = run_certify_template(cfg, mrow.community, mrow.floorplan,
                                       sheets=sheets, drive=drive)
            if ok:
                passed += 1
            else:
                failed += 1
        except Exception as e:
            print(f"  ERROR: {e}\n")
            failed += 1

        # Reload certs after each (run_certify_template saves them)
        certs = load_certifications()

    print(f"\n{'=' * 50}")
    print(f"Certify All: {passed} passed, {failed} failed, {skipped} skipped (already certified)")
    print(f"{'=' * 50}")

    all_ok = (failed == 0)
    return (all_ok, total, passed, failed)


# ── Inspect + Scan ──

def run_inspect_template(cfg: Config, community: str, floorplan: str):
    """Inspect a template: confirm invisible code, show headers, preview data."""
    sheets = SheetsClient(cfg.google.credentials_json_path, cfg.google.spreadsheet_id)
    sheets.connect()
    drive = DriveClient(cfg.google.credentials_json_path, cfg.drive.shared_drive_id)
    drive.connect()

    mapping_records = sheets.get_all_records(cfg.google.mapping_tab)
    mapping_rows = parse_mapping_tab(mapping_records)
    mrow = find_mapping_row(mapping_rows, community, floorplan)
    if not mrow:
        print(f"No MAPPING row for ({community}, {floorplan}).")
        return

    tf = drive.find_file_by_name(cfg.drive.templates_folder_id, mrow.file_name)
    if not tf:
        print(f"Template '{mrow.file_name}' not found in Drive.")
        return

    cfg.ensure_cache_dirs()
    cache_path = os.path.join(cfg.drive.download_cache_dir, mrow.file_name)
    drive.download_file(tf["id"], cache_path)

    with open(cache_path, "rb") as f:
        template_bytes = f.read()

    from docx import Document
    import io as _io
    doc = Document(_io.BytesIO(template_bytes))

    print(f"Template: {mrow.file_name}")
    print(f"Invisible code: {mrow.invisible_code}")
    print(f"Total tables in document: {len(doc.tables)}")

    try:
        match = find_table_by_invisible_code(doc, mrow.invisible_code)
        print(f"Found in table[{match.table_index}] at cell({match.cell_row},{match.cell_col})")
    except ValueError as e:
        print(f"ERROR: {e}")
        return

    from .utils import build_header_map
    header_idx = mrow.header_row - 1
    header_cells = [cell.text.strip() for cell in match.table.rows[header_idx].cells]
    hmap = build_header_map(header_cells)
    print(f"Header row {mrow.header_row}: {header_cells}")
    print(f"Header map: {hmap}")

    # Preview first 2 data rows
    data_start = header_idx + 1
    for i in range(data_start, min(data_start + 2, len(match.table.rows))):
        cells = [cell.text.strip() for cell in match.table.rows[i].cells]
        print(f"  Row {i + 1}: {cells}")


def run_scan_template(cfg: Config, file_name: str, marker_prefix: str = "[[PS|"):
    """Scan a template for marker strings."""
    drive = DriveClient(cfg.google.credentials_json_path, cfg.drive.shared_drive_id)
    drive.connect()

    tf = drive.find_file_by_name(cfg.drive.templates_folder_id, file_name)
    if not tf:
        print(f"Template '{file_name}' not found in Drive.")
        return

    cfg.ensure_cache_dirs()
    cache_path = os.path.join(cfg.drive.download_cache_dir, file_name)
    drive.download_file(tf["id"], cache_path)

    with open(cache_path, "rb") as f:
        template_bytes = f.read()

    from docx import Document
    import io as _io
    doc = Document(_io.BytesIO(template_bytes))

    results = scan_template_for_markers(doc, marker_prefix)
    print(f"Scan of '{file_name}' for marker '{marker_prefix}':")
    print(f"Total tables: {len(doc.tables)}")
    if not results:
        print("No markers found.")
    for r in results:
        print(f"  table[{r['table_index']}] cell({r['cell_row']},{r['cell_col']}): {r['snippet']}")


# ── List New Releases ──

def run_list_new_releases(cfg: Config):
    """List PDFs in the New Releases folder."""
    drive = DriveClient(cfg.google.credentials_json_path, cfg.drive.shared_drive_id)
    drive.connect()

    pdfs = drive.list_pdfs(cfg.drive.new_releases_folder_id)
    manifest = load_manifest(cfg.drive.processed_manifest)

    print(f"PDFs in New Releases: {len(pdfs)}\n")
    for pdf in pdfs:
        status = "PROCESSED" if is_already_processed(pdf, manifest) else "NEW"

        # Try release format first, then legacy
        rel_parsed = parse_release_filename(pdf["name"])
        if rel_parsed and rel_parsed.get("phase"):
            info = f" -> community={rel_parsed['community']}, phase={rel_parsed['phase']} (release PDF)"
        else:
            parsed = parse_pdf_filename(pdf["name"])
            if parsed:
                info = f" -> community={parsed['community']}, hs={parsed['homesite']}, fp={parsed['floorplan']}"
            else:
                info = " -> (cannot parse filename)"

        print(f"  [{status}] {pdf['name']} (id={pdf['id']}){info}")


# ── Sync Drive Folders ──

def run_sync_drive_folders(cfg: Config):
    """Cache folder IDs for convenience."""
    drive = DriveClient(cfg.google.credentials_json_path, cfg.drive.shared_drive_id)
    drive.connect()

    cfg.ensure_cache_dirs()
    cache = {
        "templates_folder_id": cfg.drive.templates_folder_id,
        "new_releases_folder_id": cfg.drive.new_releases_folder_id,
        "final_price_sheets_folder_id": cfg.drive.final_price_sheets_folder_id,
        "sop_folder_id": cfg.drive.sop_folder_id,
        "synced_at": datetime.now(timezone.utc).isoformat(),
    }

    # Verify each folder
    for label, fid in cache.items():
        if fid and not fid.startswith("PASTE_") and label != "synced_at":
            ok = drive.verify_folder_access(fid, label)
            cache[f"{label}_verified"] = ok

    with open(cfg.drive.folder_cache_file, "w", encoding="utf-8") as f:
        json.dump(cache, f, indent=2)

    print(f"Folder IDs cached to {cfg.drive.folder_cache_file}")
    print(json.dumps(cache, indent=2))
