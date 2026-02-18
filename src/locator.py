"""Table locator - finds the target table in a DOCX by invisible code."""

import logging
from dataclasses import dataclass
from typing import List, Optional, Tuple

from docx import Document
from docx.table import Table

logger = logging.getLogger("price_sheet_bot.locator")


@dataclass
class TableMatch:
    """Result of finding a table by invisible code."""
    table_index: int
    table: Table
    cell_row: int
    cell_col: int
    cell_text: str


def find_table_by_invisible_code(doc: Document, invisible_code: str) -> TableMatch:
    """Scan all tables/cells for the invisible_code string.

    Returns a TableMatch if found exactly once.
    Raises ValueError if not found or found in multiple tables.

    Handles broken invisible codes in templates where the closing ']]'
    may be missing (e.g. ``[[PS|COMM=NOVA|FP=02 `` instead of
    ``[[PS|COMM=NOVA|FP=02]]``).  Falls back to matching the core part
    (everything between ``[[`` and ``]]``) when an exact match fails.
    """
    if not invisible_code:
        raise ValueError("invisible_code is empty.")

    matches: List[TableMatch] = []

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if invisible_code in cell.text:
                    matches.append(TableMatch(
                        table_index=t_idx,
                        table=table,
                        cell_row=r_idx,
                        cell_col=c_idx,
                        cell_text=cell.text,
                    ))

    # Fallback: if exact match failed, try matching without the closing ']]'
    # Some templates have broken codes like "[[PS|COMM=NOVA|FP=02 " instead
    # of "[[PS|COMM=NOVA|FP=02]]"
    if len(matches) == 0 and invisible_code.endswith("]]"):
        import re
        core = invisible_code[:-2]  # e.g. "[[PS|COMM=NOVA|FP=02"
        # The character after the core must NOT be alphanumeric, to prevent
        # "FP=02" from matching "FP=02X".  It can be space, ], end-of-string, etc.
        pattern = re.escape(core) + r"(?![A-Za-z0-9])"
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if re.search(pattern, cell.text):
                        matches.append(TableMatch(
                            table_index=t_idx,
                            table=table,
                            cell_row=r_idx,
                            cell_col=c_idx,
                            cell_text=cell.text,
                        ))
        if matches:
            logger.info(
                "Invisible code '%s' not found exactly, but matched via "
                "prefix '%s' in %d cell(s).", invisible_code, core, len(matches),
            )

    if len(matches) == 0:
        raise ValueError(
            f"Invisible code '{invisible_code}' NOT FOUND in any table cell. "
            f"Scanned {len(doc.tables)} tables."
        )

    # Check if all matches are in the same table
    table_indices = set(m.table_index for m in matches)
    if len(table_indices) > 1:
        details = "; ".join(
            f"table[{m.table_index}] cell({m.cell_row},{m.cell_col})" for m in matches
        )
        raise ValueError(
            f"Invisible code '{invisible_code}' found in MULTIPLE tables: {details}. "
            f"Each invisible code must appear in exactly one table."
        )

    # Multiple matches in the same table is OK (return first)
    return matches[0]


def remove_invisible_code(table_match: TableMatch, invisible_code: str):
    """Remove the invisible_code substring from the cell where it was found.

    Also handles broken codes (missing ']]') by removing the core prefix.
    """
    cell = table_match.table.rows[table_match.cell_row].cells[table_match.cell_col]

    # Build list of strings to try removing (exact first, then prefix without ']]')
    codes_to_try = [invisible_code]
    if invisible_code.endswith("]]"):
        codes_to_try.append(invisible_code[:-2])  # prefix without closing brackets

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            for code in codes_to_try:
                if code in run.text:
                    run.text = run.text.replace(code, "")
                    # Also clean up any leftover ']]' that was in a separate run
            # Clean stray ']]' that may have been the broken closing
            if run.text.strip() == "]]":
                run.text = ""


def scan_template_for_markers(doc: Document, marker_prefix: str = "[[PS|") -> list:
    """Scan all tables/cells for a marker prefix. For debugging.

    Returns list of dicts with table_index, cell_row, cell_col, snippet.
    """
    results = []
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if marker_prefix in cell.text:
                    snippet = cell.text[:100]
                    results.append({
                        "table_index": t_idx,
                        "cell_row": r_idx,
                        "cell_col": c_idx,
                        "snippet": snippet,
                    })
    return results
